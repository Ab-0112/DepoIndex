from docx import Document

def generate_markdown(toc_items, output_path='output/toc.md'):
    with open(output_path, 'w') as f:
        f.write("| Topic | Page | Line |\n")
        f.write("|-------|------|------|\n")
        for item in toc_items:
            f.write(f"| {item['topic']} | {item['page_start']} | {item['line_start']} |\n")

def generate_docx(toc_items, output_path='output/toc.docx'):
    doc = Document()
    doc.add_heading('Deposition Topic Index', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Topic'
    hdr_cells[1].text = 'Page'
    hdr_cells[2].text = 'Line'
    
    for item in toc_items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['topic'])
        row_cells[1].text = str(item['page_start'])
        row_cells[2].text = str(item['line_start'])
    
    doc.save(output_path)
